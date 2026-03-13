"""
Generador del Manual de Usuario PEVI en formato Word (.docx)
Ejecutar: python3 generar_manual_word.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colores ──
PRIMARY = RGBColor(2, 132, 199)       # #0284c7
DARK = RGBColor(15, 23, 42)           # #0f172a
ACCENT = RGBColor(56, 189, 248)       # #38bdf8
GRAY_500 = RGBColor(100, 116, 139)
GRAY_600 = RGBColor(71, 85, 105)
WHITE = RGBColor(255, 255, 255)
SUCCESS = RGBColor(22, 163, 74)
WARNING = RGBColor(217, 119, 6)
DANGER = RGBColor(220, 38, 38)

doc = Document()

# ── Estilos de pagina ──
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Helpers ──
def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(data, col_widths=None, header_color="0f172a", header_text_color="FFFFFF"):
    """Crea una tabla con estilo profesional."""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.space_before = Pt(3)
            p.space_after = Pt(3)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

            if i == 0:  # Header row
                set_cell_shading(cell, header_color)
                run.font.color.rgb = RGBColor.from_string(header_text_color)
                run.font.bold = True
                run.font.size = Pt(8.5)
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, "f8fafc")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_heading_styled(text, level=1):
    """Agrega un heading con estilo consistente."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK if level <= 2 else PRIMARY
    return h

def add_body(text):
    """Agrega un parrafo normal."""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    p.style.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    return p

def add_tip(text, icon="ℹ️"):
    """Agrega un recuadro de consejo."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon} {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY_600
    return p

def add_bullet(text, bold_prefix=""):
    """Agrega un item de lista."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(f" {text}")
    else:
        run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_numbered(text, bold_prefix=""):
    """Agrega un item de lista numerada."""
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(f" {text}")
    else:
        run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_separator():
    """Linea separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(226, 232, 240)

# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════

# Espaciado superior
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUAL DE USUARIO")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plataforma PEVI")
run.font.size = Pt(20)
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Programa de Evaluación de Eficiencia Energética Industrial")
run.font.size = Pt(12)
run.font.color.rgb = GRAY_500

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UPME Colombia  •  MinEnergía  •  ONUDI")
run.font.size = Pt(11)
run.font.color.rgb = GRAY_600

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versión 2.0 — Febrero 2026")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_500

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("Tabla de Contenidos", level=1)
add_separator()

toc_items = [
    "1. Introducción",
    "2. Roles y Permisos",
    "3. Arquitectura del Sistema",
    "4. Inicio de Sesión",
    "5. Dashboard Principal",
    "6. Gestión de Proyectos",
    "7. Registro de Datos Energéticos",
    "8. Oportunidades de Mejora (Reducción)",
    "9. OPM - Oportunidades de Mejora Identificadas",
    "10. Documentos del Proyecto",
    "11. Generación de Informes PDF",
    "12. Gestión de Empresas",
    "13. Gestión de Usuarios",
    "14. Dashboard Estratégico (BI)",
    "15. Dashboard Nacional",
    "16. Sitio Público",
    "17. Panel de Super Administración",
    "18. Glosario de Términos",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCCION
# ══════════════════════════════════════════════════════════════

add_heading_styled("1. Introducción", level=1)
add_separator()

add_body("La Plataforma PEVI es un sistema web de gestión de auditorías energéticas desarrollado para el Programa de Eficiencia Energética Industrial (PEVI), coordinado por la UPME (Unidad de Planeación Minero Energética) de Colombia.")

add_body("El sistema permite a los centros universitarios participantes registrar, analizar y reportar los consumos energéticos de las empresas industriales auditadas, generando indicadores clave de desempeño energético, huella de carbono y oportunidades de mejora.")

add_heading_styled("Centros PEVI Participantes", level=2)
add_bullet("Universidad del Atlántico")
add_bullet("Universidad Autónoma de Occidente (UAO)")
add_bullet("Universidad de Vigo")
add_bullet("Universidad Tecnológica de Pereira (UTP)")
add_bullet("Universidad Francisco de Paula Santander (UFPS)")

add_heading_styled("Capacidades del Sistema", level=2)
add_bullet("Gestionar proyectos de auditoría energética con equipos de trabajo", "Proyectos:")
add_bullet("6 fuentes: electricidad, gas natural, carbón mineral, fuel oil, biomasa y GLP", "Registro energético:")
add_bullet("Metas de reducción y evaluación de ahorros en kWh, COP y TonCO2", "Análisis de mejoras:")
add_bullet("Reportes profesionales con KPIs, gráficos y análisis completo", "Informes PDF:")
add_bullet("Dashboards estratégico y nacional con visión consolidada", "Inteligencia de negocios:")

add_heading_styled("Acceso al Sistema", level=2)
add_styled_table([
    ["Entorno", "URL", "Descripción"],
    ["Producción", "https://pevicolombia.com", "Servidor principal (AWS)"],
    ["Aplicación interna", "/app/", "Dashboard y gestión (requiere login)"],
    ["Sitio público", "/", "Página pública con resultados agregados"],
], col_widths=[4, 6, 6])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ROLES Y PERMISOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("2. Roles y Permisos", level=1)
add_separator()

add_body("El sistema opera con un modelo de roles jerárquico que determina el nivel de acceso de cada usuario. Cada rol hereda los permisos del nivel inferior.")

add_heading_styled("Descripción de Roles", level=2)

add_styled_table([
    ["Rol", "Código", "Descripción"],
    ["Director Nacional", "DIRECTOR_NACIONAL", "Acceso total al sistema. Visión de todos los centros del país. Máximo nivel de autoridad."],
    ["Director de Centro", "DIRECTOR_CENTRO", "Gestiona su centro PEVI. Ve proyectos, equipos y métricas de su propio centro."],
    ["Profesor Líder", "PROFESOR", "Lidera proyectos de auditoría, registra empresas, gestiona equipos de trabajo."],
    ["Estudiante", "ESTUDIANTE", "Participa en equipos asignados. Registra datos energéticos en proyectos."],
], col_widths=[4, 5, 7])

add_heading_styled("Matriz de Permisos", level=2)

add_styled_table([
    ["Funcionalidad", "Estudiante", "Profesor", "Dir. Centro", "Dir. Nacional"],
    ["Ver dashboard", "✓", "✓", "✓", "✓"],
    ["Registrar datos energéticos", "✓", "✓", "✓", "✓"],
    ["Subir documentos al proyecto", "✓", "✓", "✓", "✓"],
    ["Generar informe PDF", "✓", "✓", "✓", "✓"],
    ["Crear proyectos", "✗", "✓", "✓", "✓"],
    ["Editar estructura del proyecto", "✗", "✓", "✓", "✓"],
    ["Gestionar empresas (CRUD)", "✗", "✓", "✓", "✓"],
    ["Gestionar usuarios", "✗", "✗", "✓", "✓"],
    ["Dashboard Estratégico (BI)", "✗", "✗", "✓", "✓"],
    ["Dashboard Nacional", "✗", "✗", "✗", "✓"],
    ["Panel de Super Admin", "✗", "✗", "✗", "✓ *"],
], col_widths=[5, 2.5, 2.5, 2.5, 2.5])

add_tip("* Solo usuarios con permisos de superusuario configurados en el sistema.")

add_heading_styled("Aislamiento por Centro", level=2)
add_body("Cada usuario solo puede ver y gestionar datos de su propio Centro PEVI. Las empresas, proyectos y usuarios están filtrados por centro. El Director Nacional es la única excepción, con visión completa de todo el país.")

add_tip("⚠️ Seguridad anti-escalada: Un Director de Centro no puede crear otros Directores ni Directores Nacionales. Esta restricción se aplica tanto en el formulario como en el servidor.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. ARQUITECTURA
# ══════════════════════════════════════════════════════════════

add_heading_styled("3. Arquitectura del Sistema", level=1)
add_separator()

add_heading_styled("Modelo de Entidades", level=2)
add_body("El sistema se organiza alrededor de las siguientes entidades principales:")

add_styled_table([
    ["Entidad", "Descripción", "Relaciones"],
    ["Centro PEVI", "Universidad participante en el programa", "Contiene: Usuarios, Empresas, Proyectos"],
    ["Usuario", "Persona con acceso al sistema (tiene un rol)", "Pertenece a: 1 Centro. Lidera: N Proyectos. Participa en: N Equipos"],
    ["Empresa", "Empresa industrial cliente auditada", "Pertenece a: 1 Centro. Tiene: N Proyectos"],
    ["Proyecto Auditoría", "Unidad central de trabajo", "Tiene: 1 Empresa, 1 Líder, N Equipo, 6 Fuentes Energéticas, N Documentos, N OPMs"],
    ["Fuentes Energéticas", "Electricidad, Gas Natural, Carbón, Fuel Oil, Biomasa, GLP", "Pertenece a: 1 Proyecto (relación 1:1 por tipo)"],
    ["Documento", "Archivo adjunto al proyecto", "Pertenece a: 1 Proyecto"],
    ["OPM", "Oportunidad de mejora identificada", "Pertenece a: 1 Proyecto"],
], col_widths=[3.5, 5, 7])

add_heading_styled("Ciclo de Vida del Proyecto", level=2)
add_body("Los proyectos siguen un flujo de estados definido:")

add_styled_table([
    ["Estado", "Descripción", "Transiciones Posibles"],
    ["BORRADOR", "Configuración inicial. El equipo y la empresa están siendo definidos.", "→ En Ejecución"],
    ["EN EJECUCIÓN", "Fase activa. Registro de datos energéticos en progreso.", "→ En Revisión / → Borrador"],
    ["EN REVISIÓN", "Datos completos, pendiente de validación por director.", "→ Finalizado / → En Ejecución"],
    ["FINALIZADO", "Estado terminal. Proyecto completado y cerrado.", "Ninguna (solo superadmin puede reabrir)"],
], col_widths=[3, 7, 5])

add_heading_styled("Flujo General de Trabajo", level=2)
add_body("El proceso típico de una auditoría energética sigue estos pasos:")

add_numbered("Registrar la empresa cliente en el sistema.", "Crear Empresa:")
add_numbered("Crear el proyecto vinculado a la empresa.", "Crear Proyecto:")
add_numbered("Agregar miembros al equipo de trabajo.", "Asignar Equipo:")
add_numbered("Definir la producción anual de la empresa.", "Registrar Producción:")
add_numbered("Ingresar datos de las 6 fuentes energéticas.", "Registrar Consumos:")
add_numbered("Revisar KPIs, gráficos y distribución energética.", "Analizar Resultados:")
add_numbered("Establecer porcentajes de reducción por fuente.", "Definir Metas:")
add_numbered("Registrar proyectos concretos de mejora (OPMs).", "Registrar OPMs:")
add_numbered("Descargar el informe completo en PDF.", "Generar Informe:")
add_numbered("Cambiar el estado del proyecto a Finalizado.", "Finalizar Proyecto:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. INICIO DE SESION
# ══════════════════════════════════════════════════════════════

add_heading_styled("4. Inicio de Sesión", level=1)
add_separator()

add_body("Para acceder a la plataforma, navegue a https://pevicolombia.com. El sistema lo redirigirá a la pantalla de inicio de sesión.")

add_heading_styled("Pantalla de Login", level=2)
add_body("La pantalla de inicio de sesión presenta un diseño dividido en dos paneles:")

add_bullet("Muestra los logos de los 5 centros universitarios participantes sobre un fondo con imagen de fondo energético.", "Panel izquierdo (Marca):")
add_bullet("Contiene el formulario de autenticación con campos de usuario/email y contraseña, botón de ingreso, enlace de recuperación de contraseña, y logos institucionales (MinEnergía, ONUDI, UPME).", "Panel derecho (Formulario):")

add_heading_styled("Credenciales de Acceso", level=2)
add_body("Puede iniciar sesión usando su nombre de usuario o su correo electrónico junto con su contraseña. Ambas opciones son válidas.")

add_tip("🔑 ¿Olvidó su contraseña? Use el enlace \"¿Olvidaste tu contraseña?\" en la pantalla de login para restablecerla a través de su email. Si no tiene acceso a su correo, contacte al Director de su Centro PEVI.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. DASHBOARD
# ══════════════════════════════════════════════════════════════

add_heading_styled("5. Dashboard Principal", level=1)
add_separator()
add_body("Ruta: /app/")
add_body("Al iniciar sesión, el sistema lo recibe con un saludo dinámico (Buenos días/tardes/noches según la hora) y muestra una vista adaptada a su rol.")

add_heading_styled("Vista Ejecutiva (Directores)", level=2)
add_body("Los Directores de Centro y el Director Nacional ven 4 tarjetas de indicadores clave:")

add_styled_table([
    ["Tarjeta", "Contenido", "Descripción"],
    ["Proyectos Totales", "Número entero", "Cantidad total de proyectos en su alcance"],
    ["Proyectos Activos", "Número entero", "Proyectos en estado \"En Ejecución\""],
    ["Energía Auditada", "kWh (humanizado)", "Total de energía registrada (ej: 1.2M kWh)"],
    ["Métricas & BI", "Enlace", "Acceso directo al Dashboard Estratégico o Nacional"],
], col_widths=[4, 4, 8])

add_heading_styled("Vista Operativa (Profesores y Estudiantes)", level=2)
add_body("Los profesores y estudiantes ven un panel de trabajo simplificado con acceso rápido a sus proyectos. Los profesores también ven un botón para \"Iniciar Nuevo Proyecto\".")

add_heading_styled("Tabla de Actividad Reciente", level=2)
add_body("Ambas vistas incluyen una tabla con los últimos 10 proyectos actualizados, mostrando:")

add_styled_table([
    ["Columna", "Descripción"],
    ["Proyecto / Cliente", "Nombre del proyecto y empresa auditada"],
    ["Estado", "Badge de color: Borrador (gris), En Ejecución (azul), En Revisión (amarillo), Finalizado (verde)"],
    ["Energía (kWh)", "Total de kWh auditados en el proyecto"],
    ["Última actualización", "Fecha del último cambio realizado"],
    ["Acción", "Botón para ver o gestionar el proyecto"],
], col_widths=[4, 12])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. GESTION DE PROYECTOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("6. Gestión de Proyectos", level=1)
add_separator()

add_heading_styled("6.1 Lista de Proyectos", level=2)
add_body("Ruta: /app/proyectos/")
add_body("Muestra todos los proyectos según su rol:")

add_bullet("\"Portafolio Nacional de Proyectos\" — Ve todos los centros del país", "Director Nacional:")
add_bullet("Ve los proyectos de su centro", "Director de Centro:")
add_bullet("\"Mis Proyectos Liderados\" + proyectos donde participa como equipo", "Profesor:")
add_bullet("Solo proyectos asignados a su equipo", "Estudiante:")

add_heading_styled("Filtros Disponibles", level=3)
add_styled_table([
    ["Filtro", "Opciones", "Disponible para"],
    ["Búsqueda", "Texto libre (nombre de proyecto o empresa)", "Todos los roles"],
    ["Estado", "Borrador, Ejecución, Revisión, Finalizado", "Todos los roles"],
    ["Fase", "Fase 1, Fase 2, Fase 3, Fase 4", "Todos los roles"],
    ["Líder", "Lista de profesores + opción \"Mis Proyectos\"", "Solo directores"],
    ["Centro PEVI", "Lista de centros activos", "Solo Director Nacional"],
], col_widths=[3, 6, 4])

add_heading_styled("Columnas de la Tabla", level=3)
add_styled_table([
    ["Columna", "Contenido"],
    ["Proyecto / Cliente", "Nombre del proyecto (enlace) + razón social de la empresa"],
    ["Líder y Equipo", "Nombre del líder (★ \"Yo\" si es usted) + badge con cantidad de miembros"],
    ["Centro PEVI", "Nombre del centro (solo visible para Director Nacional)"],
    ["Estado", "Badge de color según el estado del proyecto"],
    ["Acciones", "Botón \"Ver\" (directores) o \"Gestionar\" (operativos)"],
], col_widths=[4, 12])

add_heading_styled("6.2 Crear un Proyecto Nuevo", level=2)
add_body("Ruta: /app/proyectos/nuevo/")
add_body("Disponible para: Profesor, Director de Centro, Director Nacional.")

add_numbered("Ingrese un nombre descriptivo. Ejemplo: \"Evaluación Energética Planta 2 - 2025\".", "Nombre de la auditoría:")
add_numbered("Seleccione la empresa a auditar del menú desplegable. Si no existe, use el enlace \"Regístrala aquí\" para crearla primero.", "Empresa cliente:")
add_numbered("Si usted es Profesor, se pre-selecciona automáticamente. Los directores pueden asignar a otro líder.", "Líder del proyecto:")
add_numbered("Use el selector interactivo por roles. Haga clic en la pill del rol deseado (Estudiante, Profesor, etc.) para ver las personas disponibles, luego haga clic en sus tarjetas para agregarlas al equipo.", "Equipo de trabajo:")
add_numbered("Seleccione la fase PEVI correspondiente si aplica (Fase 1 a 4).", "Fase del programa (opcional):")
add_numbered("Defina las fechas de inicio y cierre estimado del proyecto.", "Cronograma:")
add_numbered("Haga clic en el botón \"Crear Proyecto\".", "Guardar:")

add_tip("💡 Selector de equipo por roles: El sistema agrupa a los usuarios por su rol. Los miembros seleccionados aparecen como badges de colores con botón X para quitarlos. Puede buscar por nombre dentro de cada rol.")

add_heading_styled("6.3 Editar un Proyecto", level=2)
add_body("Ruta: /app/proyectos/<id>/editar/")
add_body("Solo el líder del proyecto, el director de su centro, o el Director Nacional pueden editar la estructura del proyecto.")

add_tip("⚠️ Transferencia de liderazgo: Si un profesor cambia el líder del proyecto a otra persona, el sistema solicita doble confirmación porque perderá acceso de edición al proyecto.")

add_heading_styled("6.4 Estados y Transiciones", level=2)

add_styled_table([
    ["Estado Actual", "Acción", "Estado Resultante", "Quién puede hacerlo"],
    ["Borrador", "Iniciar Ejecución", "En Ejecución", "Líder, Director"],
    ["En Ejecución", "Enviar a Revisión", "En Revisión", "Líder, Director"],
    ["En Ejecución", "Devolver a Borrador", "Borrador", "Líder, Director"],
    ["En Revisión", "Finalizar Proyecto", "Finalizado", "Director"],
    ["En Revisión", "Devolver a Ejecución", "En Ejecución", "Director"],
    ["Finalizado", "Reabrir (excepcional)", "Según selección", "Solo Superadmin"],
], col_widths=[3, 3.5, 3.5, 4])

add_heading_styled("6.5 Expediente del Proyecto (Vista Detalle)", level=2)
add_body("Ruta: /app/proyectos/<id>/")
add_body("Esta es la pantalla central del proyecto. Contiene toda la información organizada en secciones:")

add_heading_styled("Barra de Estado Superior", level=3)
add_body("Muestra el estado actual del proyecto con badge de color, la fase (si aplica), fecha de inicio y botones de acción para cambiar estado, editar estructura o generar el informe PDF.")

add_heading_styled("3 KPI Principales", level=3)
add_styled_table([
    ["KPI", "Unidad", "Descripción"],
    ["Consumo Total", "kWh equivalentes", "Suma de energía eléctrica y térmica convertida"],
    ["Huella de Carbono", "TonCO2/año", "Total de emisiones de CO2 de todas las fuentes"],
    ["Costo Operativo", "COP/año", "Costo total anual de energía"],
], col_widths=[4, 4, 8])

add_heading_styled("Bitácora de Registros", level=3)
add_body("7 tarjetas interactivas que muestran el estado de cada registro:")
add_bullet("Producción/Actividad: Cantidad producida + unidad + indicador IDES")
add_bullet("Electricidad: Consumo kWh + Costo COP")
add_bullet("Gas Natural: Volumen m³ + Equivalente kWh")
add_bullet("Carbón Mineral: Masa Ton + Emisiones TonCO2")
add_bullet("Fuel Oil: Volumen Gal + Energía kWh")
add_bullet("Biomasa: Tipo + Masa Ton")
add_bullet("Gas Propano (GLP): Masa kg + Energía kWh")

add_heading_styled("4 Gráficos de Análisis", level=3)
add_styled_table([
    ["Gráfico", "Tipo", "Descripción"],
    ["Matriz de Consumo (kWh)", "Dona", "Distribución porcentual de energía por fuente"],
    ["Huella de Carbono (TonCO2)", "Dona", "Distribución de emisiones por fuente"],
    ["Balance Térmico (MBTU)", "Barras", "Comparación eléctrica vs térmica"],
    ["Distribución de Costos (COP)", "Barras", "Costos desglosados por fuente energética"],
], col_widths=[5, 3, 8])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. REGISTRO ENERGETICO
# ══════════════════════════════════════════════════════════════

add_heading_styled("7. Registro de Datos Energéticos", level=1)
add_separator()

add_body("Desde el expediente del proyecto, haga clic en la tarjeta de la fuente energética que desea registrar. El sistema presenta un formulario estandarizado.")
add_body("Ruta: /app/proyectos/<id>/registro/<tipo_energia>/")

add_heading_styled("7.1 Fuentes Energéticas Disponibles", level=2)

add_styled_table([
    ["Fuente", "Unidad de Consumo", "Poder Calorífico", "Factor de Emisión"],
    ["Electricidad", "kWh", "N/A (ya es energía)", "kgCO2/kWh"],
    ["Gas Natural", "m³", "kJ/m³", "kgCO2/m³"],
    ["Carbón Mineral", "Toneladas", "MJ/kg", "kgCO2/Ton"],
    ["Fuel Oil", "Galones", "MJ/usgal", "kgCO2/Gal"],
    ["Biomasa", "Toneladas", "kJ/kg", "kgCO2/Ton"],
    ["Gas Propano (GLP)", "kg", "MJ/kg", "kgCO2/kg"],
], col_widths=[4, 3.5, 3.5, 3.5])

add_heading_styled("7.2 Estructura del Formulario", level=2)
add_body("Cada formulario de registro contiene las siguientes secciones:")

add_numbered("Seleccione el tipo: Bagazo, Cascarilla de arroz, Madera, etc.", "Tipo de biomasa (solo Biomasa):")
add_numbered("Ingrese el consumo anual en las unidades originales de la fuente.", "Datos de consumo:")
add_numbered("Registre el costo total anual (COP) y opcionalmente el costo por unidad.", "Información de costos:")
add_numbered("Ingrese el poder calorífico en la unidad correspondiente. Use el botón de conversor de unidades si necesita convertir.", "Parámetros técnicos (excepto electricidad):")
add_numbered("El sistema calcula automáticamente el consumo equivalente en kWh. Este campo es de solo lectura.", "Conversión energética (automática):")
add_numbered("Ingrese el factor de emisión para calcular las emisiones de CO2. Puede consultar los factores oficiales de UPME.", "Impacto ambiental:")

add_heading_styled("7.3 Conversión Automática a kWh", level=2)
add_body("El sistema convierte todos los combustibles a kWh equivalente usando la siguiente fórmula:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Energía (kWh) = (Consumo × Factor_Unidad × Poder_Calorífico × Factor_Energía) / 3,600")
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = PRIMARY

add_heading_styled("Factores de Conversión por Energético", level=3)
add_styled_table([
    ["Energético", "factor_unidad", "factor_energía", "Explicación"],
    ["Gas Natural", "1.0", "1.0", "PC ya viene en kJ/m³, no requiere conversión"],
    ["Carbón Mineral", "1000 (Ton→kg)", "1000 (MJ→kJ)", "Convierte toneladas a kg y MJ a kJ"],
    ["Fuel Oil", "1.0", "1000 (MJ→kJ)", "PC en MJ/usgal, convierte MJ a kJ"],
    ["Biomasa", "1000 (Ton→kg)", "1.0", "Convierte toneladas a kg, PC ya en kJ/kg"],
    ["Gas Propano", "1.0", "1000 (MJ→kJ)", "PC en MJ/kg, convierte MJ a kJ"],
], col_widths=[3, 3, 3, 6])

add_heading_styled("7.4 Cálculo de Emisiones", level=2)
add_body("Las emisiones se calculan sobre el consumo en unidades originales, NO sobre kWh equivalentes:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Emisiones (TonCO2) = (Consumo_Anual_Original × Factor_Emisión) / 1,000")
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = SUCCESS

add_tip("💡 Formato numérico: Puede ingresar números con comas (1,200.50) o sin ellas (1200.50). El sistema las procesa automáticamente.")

add_heading_styled("7.5 Registrar Producción", level=2)
add_body("Ruta: /app/proyectos/<id>/registro/produccion/")
add_body("Antes de registrar consumos, defina la producción de la empresa:")
add_bullet("Volumen de producción anual", "Cantidad producida (Año Base):")
add_bullet("Ej: Toneladas, Unidades, m³", "Unidad de medida:")
add_body("Con estos datos, el sistema calcula el IDES (Indicador de Desempeño Energético Específico) = Energía Total / Producción Total.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. OPORTUNIDADES DE MEJORA
# ══════════════════════════════════════════════════════════════

add_heading_styled("8. Oportunidades de Mejora (Reducción)", level=1)
add_separator()

add_body("Desde el expediente del proyecto, la sección \"Oportunidades de Mejora\" permite establecer metas de reducción porcentual para cada fuente energética y calcular los ahorros potenciales.")

add_heading_styled("Panel de Comparación", level=2)
add_body("Un panel visual muestra lado a lado el \"Consumo Actual\" y el \"Consumo Proyectado\" con una flecha indicando la reducción estimada.")

add_heading_styled("3 KPI de Ahorro", level=2)
add_styled_table([
    ["KPI", "Unidad", "Cálculo"],
    ["Ahorro Energía", "kWh", "Σ(kWh × % reducción) / 100"],
    ["Ahorro Costo", "COP", "Σ(Costo × % reducción) / 100"],
    ["CO2 Evitado", "TonCO2", "Σ(Emisiones × % reducción) / 100"],
], col_widths=[4, 3, 8])

add_heading_styled("Cómo Definir Metas de Reducción", level=2)
add_numbered("En la tabla de oportunidades, localice la fila de la fuente energética deseada.")
add_numbered("Haga clic en el botón \"Editar\" de esa fila.")
add_numbered("En el modal que aparece, ingrese el porcentaje de reducción (0% a 100%).")
add_numbered("El sistema calcula automáticamente los ahorros en kWh, COP y TonCO2.")
add_numbered("Los 3 KPI boxes y el gráfico de barras horizontales se actualizan con los nuevos totales.")

add_heading_styled("Tabla Detallada", level=2)
add_body("La tabla muestra por cada fuente energética:")
add_styled_table([
    ["Columna", "Descripción"],
    ["Energético", "Nombre de la fuente de energía"],
    ["Consumo Actual", "kWh equivalentes actuales"],
    ["Costo Actual", "COP anuales actuales"],
    ["% Reducción", "Meta de reducción definida por el usuario"],
    ["Ahorro Energía", "kWh que se ahorrarían con la reducción"],
    ["Ahorro Costo", "COP que se ahorrarían"],
    ["CO2 Evitado", "TonCO2 que se evitarían"],
    ["Acciones", "Botón para editar el porcentaje de reducción"],
], col_widths=[4, 12])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. OPM IDENTIFICADAS
# ══════════════════════════════════════════════════════════════

add_heading_styled("9. OPM - Oportunidades de Mejora Identificadas", level=1)
add_separator()

add_body("Las OPM son proyectos específicos de mejora identificados durante la auditoría. Se registran con métricas técnicas y financieras para evaluar su viabilidad.")

add_heading_styled("Campos de Cada OPM", level=2)
add_styled_table([
    ["Campo", "Tipo", "Descripción"],
    ["Código", "Texto (20 car.)", "Identificador único. Ej: OPM-01"],
    ["Descripción", "Texto (300 car.)", "Descripción de la mejora propuesta"],
    ["Energético", "Selección", "Electricidad, Gas Natural, Carbón, Fuel Oil, Biomasa, GLP u Otro"],
    ["Ahorro Energía", "Numérico", "kWh/año ahorrados"],
    ["Costos Evitados", "Numérico", "MCOP/año de ahorro"],
    ["Emisiones Evitadas", "Numérico", "TonCO2/año evitadas"],
    ["Inversión", "Numérico", "COP requeridos para implementar"],
    ["VPN", "Numérico", "Valor Presente Neto (COP)"],
    ["TIR", "Porcentaje", "Tasa Interna de Retorno (%)"],
    ["Payback", "Numérico", "Período de recuperación (meses)"],
    ["Observaciones", "Texto largo", "Notas adicionales (opcional)"],
], col_widths=[3.5, 3, 8])

add_heading_styled("Flujo de Registro", level=2)
add_numbered("En el expediente del proyecto, desplácese a la sección \"Oportunidades de Mejora - Proyectos Identificados\".")
add_numbered("Haga clic en \"Nueva Oportunidad\".")
add_numbered("Complete el formulario con los datos técnicos y financieros.")
add_numbered("Guarde. La OPM aparecerá en la tabla con totales automáticos en el pie.")
add_body("Puede editar o eliminar OPMs existentes usando los botones de acción en cada fila. La eliminación requiere confirmación.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. DOCUMENTOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("10. Documentos del Proyecto", level=1)
add_separator()

add_body("Cada proyecto puede tener múltiples documentos adjuntos (informes, fotografías, hojas de cálculo, etc.).")

add_heading_styled("Subir un Documento", level=2)
add_numbered("En el expediente del proyecto, desplácese a la sección \"Documentos\".")
add_numbered("Haga clic en \"Subir Documento\".")
add_numbered("En el modal, ingrese una descripción y seleccione el archivo.")
add_numbered("Formatos soportados: PDF, Excel, Word, imágenes. Tamaño máximo: 10 MB.")
add_numbered("Haga clic en \"Guardar\". El documento aparecerá en la tabla.")

add_heading_styled("Tabla de Documentos", level=2)
add_styled_table([
    ["Columna", "Descripción"],
    ["Nombre del Archivo", "Descripción ingresada al subir"],
    ["Fecha de Carga", "Fecha y hora de la subida"],
    ["Tamaño", "Peso del archivo"],
    ["Descargar", "Botón para descargar el archivo original"],
], col_widths=[4, 12])

# ══════════════════════════════════════════════════════════════
# 11. INFORMES PDF
# ══════════════════════════════════════════════════════════════

add_heading_styled("11. Generación de Informes PDF", level=1)
add_separator()
add_body("Ruta: /app/proyectos/<id>/informe/pdf/")
add_body("Desde el expediente del proyecto, haga clic en el botón \"Informe PDF\" en la barra de acciones superior.")

add_heading_styled("Contenido del Informe", level=2)
add_numbered("Nombre del proyecto, empresa, centro PEVI, fecha de generación.", "Encabezado:")
add_numbered("KPIs principales: Energía Total (kWh), Emisiones (TonCO2), Costo Operativo (COP), Indicador IDES.", "Resumen Ejecutivo:")
add_numbered("Todas las fuentes energéticas con consumo original, equivalente kWh, emisiones y costos.", "Tabla de Consumos:")
add_numbered("Gráfico de barras horizontales por fuente con porcentajes.", "Distribución de Consumo:")
add_numbered("Comparación eléctrica vs térmica en MBTU, costo por kWh, emisiones por producción.", "Balance Térmico:")
add_numbered("Ahorros proyectados en kWh, COP y TonCO2, tabla de OPMs con métricas financieras.", "Oportunidades de Mejora:")
add_numbered("Lista de miembros del equipo del proyecto.", "Equipo de Trabajo:")

add_body("El archivo se descarga automáticamente con el formato: Informe_PEVI_[Empresa]_[ID].pdf")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. GESTION DE EMPRESAS
# ══════════════════════════════════════════════════════════════

add_heading_styled("12. Gestión de Empresas", level=1)
add_separator()
add_body("Ruta: /app/empresas/")
add_body("Disponible para: Profesor, Director de Centro, Director Nacional.")

add_heading_styled("Directorio de Empresas", level=2)
add_styled_table([
    ["Columna", "Descripción"],
    ["Razón Social / NIT", "Nombre legal y número de identificación tributaria"],
    ["Sector Productivo", "Badge con el sector industrial"],
    ["Ubicación", "Ciudad con ícono de geolocalización"],
    ["Contacto", "Nombre y email de la persona de contacto"],
    ["Acciones", "Botón de edición"],
], col_widths=[4, 12])

add_heading_styled("Registrar una Empresa Nueva", level=2)
add_body("Ruta: /app/empresas/nueva/")
add_numbered("Se asigna automáticamente según su centro (Director Nacional puede elegir).", "Centro PEVI:")
add_numbered("Nombre legal completo de la empresa.", "Razón Social:")
add_numbered("Número de identificación tributaria.", "NIT:")
add_numbered("Sector industrial de la empresa.", "Sector Productivo:")
add_numbered("Ubicación física de la empresa.", "Dirección y Ciudad:")
add_numbered("Nombre, email y teléfono de la persona de contacto.", "Datos de contacto:")

add_tip("💡 Puede crear una empresa directamente desde el formulario de creación de proyecto usando el enlace \"Regístrala aquí\".")

# ══════════════════════════════════════════════════════════════
# 13. GESTION DE USUARIOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("13. Gestión de Usuarios", level=1)
add_separator()
add_body("Ruta: /app/equipo/")
add_body("Disponible solo para: Director de Centro, Director Nacional.")

add_heading_styled("Lista de Equipo", level=2)
add_styled_table([
    ["Columna", "Descripción"],
    ["Usuario", "Avatar con iniciales, nombre completo y email"],
    ["Rol / Cargo", "Badge de rol + cargo opcional"],
    ["Estado", "Activo (verde) o Inactivo (rojo)"],
    ["Último acceso", "Fecha del último login o \"Nunca\""],
    ["Acciones", "Menú: Editar, Eliminar (con confirmación)"],
], col_widths=[4, 12])

add_heading_styled("Crear un Usuario", level=2)
add_body("Ruta: /app/equipo/nuevo/")
add_numbered("Nombre de usuario único para el login.", "Username:")
add_numbered("Correo electrónico (también puede usarse para login).", "Email:")
add_numbered("Datos personales del usuario.", "Nombre y Apellido:")
add_numbered("Seleccione: Estudiante, Profesor o Director de Centro.", "Rol:")
add_numbered("Se asigna automáticamente según su centro.", "Centro PEVI:")
add_numbered("Posición o título (opcional).", "Cargo:")
add_numbered("Marque para habilitar el acceso al sistema.", "Activo:")

add_tip("⚠️ Seguridad: Un Director de Centro NO puede crear otros Directores de Centro ni Directores Nacionales. Esta restricción previene escalada de privilegios.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. DASHBOARD ESTRATEGICO
# ══════════════════════════════════════════════════════════════

add_heading_styled("14. Dashboard Estratégico (BI)", level=1)
add_separator()
add_body("Ruta: /app/metricas/estrategico/")
add_body("Disponible para: Director de Centro, Director Nacional.")

add_heading_styled("Alcance de Datos", level=2)
add_bullet("Ve solo los proyectos de su centro", "Director de Centro:")
add_bullet("Ve todos los proyectos del país", "Director Nacional / Superadmin:")

add_heading_styled("Filtros", level=2)
add_styled_table([
    ["Filtro", "Descripción"],
    ["Proyecto", "Seleccionar un proyecto específico para ver su detalle"],
    ["Líder", "Filtrar por profesor líder del proyecto"],
], col_widths=[4, 12])

add_heading_styled("Indicadores y Gráficos", level=2)
add_body("Tarjetas KPI superiores:")
add_bullet("Total de Proyectos")
add_bullet("Energía Total (kWh eléctricos + térmicos)")
add_bullet("Costo Total (COP)")
add_bullet("Emisiones Totales (TonCO2)")
add_bullet("Desglose: Eléctrica (kWh) vs Térmica (MBTU)")

add_body("Gráficos:")
add_bullet("Energía por fuente (gráfico de dona)")
add_bullet("Costos por fuente (gráfico de dona)")
add_bullet("Balance térmico (gráfico de barras: eléctrico vs térmico en MBTU)")

add_heading_styled("Tabla Detallada de Proyectos", level=2)
add_body("Incluye todas las métricas por proyecto: Empresa, Líder, Producción, Unidad, Energía (Total/Eléctrica/Térmica), Costo, Emisiones e IDES.")

# ══════════════════════════════════════════════════════════════
# 15. DASHBOARD NACIONAL
# ══════════════════════════════════════════════════════════════

add_heading_styled("15. Dashboard Nacional", level=1)
add_separator()
add_body("Ruta: /app/metricas/nacional/")
add_body("Disponible exclusivamente para: Director Nacional.")

add_heading_styled("Modo A: Comparación Nacional (por defecto)", level=2)
add_body("Muestra una tabla comparativa de todos los centros PEVI:")

add_styled_table([
    ["Columna", "Descripción"],
    ["Centro", "Nombre del centro PEVI"],
    ["Región", "Ubicación geográfica"],
    ["Proyectos", "Cantidad de proyectos"],
    ["Energía (kWh)", "Total de energía auditada"],
    ["Emisiones", "Total TonCO2"],
    ["Promedio kWh/Proyecto", "Indicador de intensidad por proyecto"],
], col_widths=[4, 12])

add_body("El ranking está ordenado por energía (descendente). KPIs agregados: Total Centros, Total Proyectos, Total Energía, Total Emisiones.")

add_heading_styled("Modo B: Detalle de Centro", level=2)
add_body("Seleccione un centro en el dropdown para ver su análisis detallado con los mismos KPIs y gráficos del Dashboard Estratégico, pero filtrados a ese centro específico.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 16. SITIO PUBLICO
# ══════════════════════════════════════════════════════════════

add_heading_styled("16. Sitio Público", level=1)
add_separator()
add_body("Páginas de acceso público sin necesidad de autenticación.")

add_styled_table([
    ["Página", "Ruta", "Contenido"],
    ["Inicio", "/", "Landing page con KPIs públicos (proyectos finalizados, centros activos) y noticias recientes"],
    ["Nosotros", "/nosotros/", "Información sobre el programa PEVI y sus objetivos"],
    ["Centros", "/centros/", "Directorio de centros con modals que muestran gráficos de proyectos, sectores y tendencias"],
    ["Biblioteca", "/biblioteca/", "Guías técnicas de UPME: Bombeo, Motores, Vapor, etc. Con enlaces a documentos PDF"],
    ["Resultados", "/resultados/", "Resultados agregados y anonimizados con filtros, gráficos y exportación"],
], col_widths=[2.5, 3, 10])

add_heading_styled("Página de Resultados (/resultados/)", level=2)
add_body("Permite consultar resultados agregados del programa con filtros por año, región/centro y sector industrial.")

add_heading_styled("KPIs Públicos", level=3)
add_body("Total Proyectos, Total Centros, Total Empresas, Sectores Únicos, Energía Total (MWh), Emisiones Totales (TonCO2).")

add_heading_styled("Visualizaciones", level=3)
add_bullet("Energía por fuente (gráfico apilado)")
add_bullet("Proyectos en el tiempo (línea, 6 años)")
add_bullet("Proyectos por sector (barras, top 8)")
add_bullet("Proyectos por región (dona)")
add_bullet("Proyectos por centro (barras, top 10)")

add_tip("🔒 Todos los datos públicos son agregados y anonimizados. No se expone información específica de empresas ni proyectos individuales.")

# ══════════════════════════════════════════════════════════════
# 17. PANEL ADMIN
# ══════════════════════════════════════════════════════════════

add_heading_styled("17. Panel de Super Administración", level=1)
add_separator()
add_body("Ruta: /app/control/")
add_body("Acceso exclusivo para cuentas con permisos de superusuario.")

add_styled_table([
    ["Módulo", "Ruta", "Acciones"],
    ["Centros PEVI", "/app/control/centros/", "Crear, editar y eliminar centros PEVI"],
    ["Usuarios", "/app/control/usuarios/", "Gestionar cualquier usuario sin restricciones de centro"],
    ["Noticias", "/app/control/noticias/", "Crear, editar y eliminar noticias del sitio público"],
], col_widths=[3, 5, 7])

add_body("También disponible: Django Admin en /admin/, que proporciona acceso directo a todos los modelos del sistema, logs de actividad y configuraciones avanzadas.")

add_tip("⚠️ Precaución: Las acciones en el panel de administración no tienen las mismas validaciones de seguridad que la interfaz principal. Use con cuidado.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 18. GLOSARIO
# ══════════════════════════════════════════════════════════════

add_heading_styled("18. Glosario de Términos", level=1)
add_separator()

add_styled_table([
    ["Término", "Definición"],
    ["PEVI", "Programa de Evaluación de Eficiencia Energética Industrial. Programa nacional coordinado por UPME."],
    ["UPME", "Unidad de Planeación Minero Energética. Entidad gubernamental colombiana que coordina el programa."],
    ["Centro PEVI", "Universidad participante en el programa. Cada centro gestiona sus propios proyectos."],
    ["kWh", "Kilovatio-hora. Unidad estándar de energía utilizada como denominador común."],
    ["MBTU", "Miles de BTU (British Thermal Units). Unidad de energía térmica. 1 MBTU = 293.07 kWh."],
    ["TonCO2", "Toneladas de dióxido de carbono. Medida de emisiones de gases de efecto invernadero."],
    ["IDES", "Indicador de Desempeño Energético Específico. Razón: Energía Total / Producción Total."],
    ["Poder Calorífico", "Cantidad de energía liberada al quemar una unidad de combustible (PC Superior o Inferior)."],
    ["Factor de Emisión", "Cantidad de CO2 emitida por unidad de combustible consumido (kgCO2/unidad)."],
    ["OPM", "Oportunidad de Mejora. Proyecto concreto de mejora energética con análisis técnico-financiero."],
    ["VPN", "Valor Presente Neto. Indicador financiero que mide la rentabilidad de una inversión."],
    ["TIR", "Tasa Interna de Retorno. Porcentaje de retorno esperado de una inversión."],
    ["Payback", "Período de Recuperación. Tiempo estimado para recuperar la inversión (meses)."],
    ["GLP", "Gas Licuado de Petróleo. También conocido como Gas Propano."],
    ["COP", "Pesos colombianos. Moneda oficial de Colombia utilizada en todos los cálculos de costos."],
    ["NIT", "Número de Identificación Tributaria. Identificador fiscal único de empresas colombianas."],
    ["CRUD", "Create, Read, Update, Delete. Operaciones básicas de gestión de datos."],
], col_widths=[3.5, 12])

# ══════════════════════════════════════════════════════════════
# CONTRAPORTADA
# ══════════════════════════════════════════════════════════════

doc.add_page_break()

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plataforma PEVI")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Programa de Evaluación de Eficiencia Energética Industrial")
run.font.size = Pt(11)
run.font.color.rgb = GRAY_500

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UPME Colombia  •  MinEnergía  •  ONUDI")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_600

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("© 2026 — Todos los derechos reservados")
run.font.size = Pt(9)
run.font.color.rgb = GRAY_500

# ── Guardar ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Manual_Usuario_PEVI.docx")
doc.save(output_path)
print(f"✓ Manual generado exitosamente: {output_path}")
print(f"  Tamaño: {os.path.getsize(output_path) / 1024:.0f} KB")
