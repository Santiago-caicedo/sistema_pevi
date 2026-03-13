"""
Generador del Informe de Impacto y Potencia del Software PEVI
Documento tipo showcase/brochure con espacios para capturas de pantalla.
Ejecutar: python3 generar_informe_impacto.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colores ──
PRIMARY = RGBColor(2, 132, 199)
PRIMARY_DARK = RGBColor(3, 105, 161)
DARK = RGBColor(15, 23, 42)
DARK_SOFT = RGBColor(30, 41, 59)
ACCENT = RGBColor(56, 189, 248)
GRAY_400 = RGBColor(148, 163, 184)
GRAY_500 = RGBColor(100, 116, 139)
GRAY_600 = RGBColor(71, 85, 105)
WHITE = RGBColor(255, 255, 255)
SUCCESS = RGBColor(22, 163, 74)
WARNING = RGBColor(217, 119, 6)
DANGER = RGBColor(220, 38, 38)
INFO = RGBColor(8, 145, 178)

doc = Document()

# ── Estilos de pagina ──
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Helpers ──
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="d1d5db", width="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{width}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def add_screenshot_placeholder(label, height_cm=8):
    """Agrega un recuadro gris con instrucciones para insertar captura."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "f1f5f9")
    set_cell_borders(cell, color="cbd5e1", width="6")
    cell.height = Cm(height_cm)

    # Icono de imagen
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.space_before = Pt(max(10, (height_cm * 8) - 30))
    run1 = p1.add_run("📷")
    run1.font.size = Pt(28)

    # Texto de instruccion
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"INSERTAR CAPTURA: {label}")
    run2.font.size = Pt(10)
    run2.font.bold = True
    run2.font.color.rgb = GRAY_500

    # Instruccion
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Clic derecho → Insertar imagen → Seleccionar archivo")
    run3.font.size = Pt(8)
    run3.font.color.rgb = GRAY_400
    run3.font.italic = True

    doc.add_paragraph()  # Espaciado

def add_two_screenshots(label1, label2, height_cm=7):
    """Dos placeholders lado a lado."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    for i, label in enumerate([label1, label2]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "f1f5f9")
        set_cell_borders(cell, color="cbd5e1", width="6")
        cell.height = Cm(height_cm)

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.space_before = Pt(max(10, (height_cm * 8) - 30))
        run1 = p1.add_run("📷")
        run1.font.size = Pt(22)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8.5)
        run2.font.bold = True
        run2.font.color.rgb = GRAY_500

    doc.add_paragraph()

def add_section_header(number, title, subtitle=""):
    """Header de seccion con numero y titulo grande."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run_num = p.add_run(f"0{number}  " if number < 10 else f"{number}  ")
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = PRIMARY
    run_num.font.bold = True
    run_title = p.add_run(title.upper())
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = DARK
    run_title.font.bold = True

    if subtitle:
        p2 = doc.add_paragraph()
        run_sub = p2.add_run(subtitle)
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = GRAY_500
        run_sub.font.italic = True

    # Linea azul
    p_line = doc.add_paragraph()
    run_line = p_line.add_run("━" * 60)
    run_line.font.size = Pt(6)
    run_line.font.color.rgb = PRIMARY

def add_body(text, bold=False, color=None, size=10, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_body_rich(parts):
    """parts = [(text, bold, color, size), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold, color, size in parts:
        run = p.add_run(text)
        run.font.size = Pt(size or 10)
        run.font.name = 'Calibri'
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
    return p

def add_kpi_row(kpis):
    """kpis = [(valor, label, color_hex), ...]"""
    table = doc.add_table(rows=1, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    for i, (valor, label, color_hex) in enumerate(kpis):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, color_hex)
        set_cell_borders(cell, color=color_hex.replace("#", ""), width="4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.space_before = Pt(12)
        run1 = p1.add_run(valor)
        run1.font.size = Pt(20)
        run1.font.bold = True
        run1.font.color.rgb = DARK

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.space_after = Pt(12)
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.bold = True
        run2.font.color.rgb = GRAY_600
        run2.font.all_caps = True

    doc.add_paragraph()

def add_impact_card(icon, title, description, color=PRIMARY):
    """Tarjeta de impacto con icono."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon}  {title}")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(description)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = GRAY_600

def add_styled_table(data, col_widths=None, header_color="0f172a"):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

            if i == 0:
                set_cell_shading(cell, header_color)
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(8.5)
            elif i % 2 == 0:
                set_cell_shading(cell, "f8fafc")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
        r = p.add_run(f" {text}")
    else:
        r = p.add_run(text)
    r.font.size = Pt(10)

def add_quote(text):
    """Cita destacada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    run = p.add_run(f"« {text} »")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = PRIMARY_DARK

# ══════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PLATAFORMA PEVI")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━")
run.font.size = Pt(12)
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Informe de Capacidades e Impacto")
run.font.size = Pt(18)
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema Integral de Gestión de Auditorías Energéticas Industriales")
run.font.size = Pt(11)
run.font.color.rgb = GRAY_500

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Programa de Evaluación de Eficiencia Energética Industrial")
run.font.size = Pt(11)
run.font.color.rgb = GRAY_600

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UPME  •  Ministerio de Energía  •  ONUDI")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_500

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Colombia, 2026")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_400

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════

add_section_header(1, "Resumen Ejecutivo", "La plataforma que transforma la gestión energética industrial en Colombia")

add_body("La Plataforma PEVI es un sistema web de última generación diseñado para gestionar, analizar y reportar auditorías energéticas industriales a escala nacional. Desarrollada para el Programa de Eficiencia Energética Industrial (PEVI) de la UPME, la plataforma conecta una red de 5 centros universitarios que realizan diagnósticos energéticos a empresas de todo el país.")

add_quote("Una sola plataforma. 5 universidades. Todo el país. Datos energéticos en tiempo real.")

add_body("El sistema no es una simple base de datos: es un ecosistema completo que automatiza conversiones energéticas, calcula emisiones de CO2, genera informes profesionales, y ofrece dashboards de inteligencia de negocios para la toma de decisiones a nivel estratégico y nacional.", size=10)

add_kpi_row([
    ("5", "CENTROS\nUNIVERSITARIOS", "e0f2fe"),
    ("6", "FUENTES\nENERGÉTICAS", "fef3c7"),
    ("4", "DASHBOARDS\nINTERACTIVOS", "dcfce7"),
    ("1", "INFORME PDF\nAUTOMÁTICO", "fee2e2"),
])

add_screenshot_placeholder("Vista general de la Plataforma — Página de Inicio / Landing Page", height_cm=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  IMPACTO Y ALCANCE
# ══════════════════════════════════════════════════════════════════

add_section_header(2, "Impacto y Alcance Nacional", "Una red universitaria conectada para la eficiencia energética de Colombia")

add_body("La Plataforma PEVI materializa la visión de un sistema nacional unificado donde múltiples instituciones académicas colaboran en tiempo real para mejorar la eficiencia energética del sector industrial colombiano.")

add_heading = doc.add_heading("Red de Centros PEVI", level=2)
for run in add_heading.runs:
    run.font.color.rgb = DARK

add_styled_table([
    ["Centro PEVI", "Ubicación", "Cobertura Regional"],
    ["Universidad del Atlántico", "Barranquilla", "Costa Caribe"],
    ["Universidad Autónoma de Occidente (UAO)", "Cali", "Valle del Cauca y Suroccidente"],
    ["Universidad de Vigo", "Vigo, España", "Cooperación Internacional"],
    ["Universidad Tecnológica de Pereira (UTP)", "Pereira", "Eje Cafetero"],
    ["Universidad Francisco de Paula Santander (UFPS)", "Cúcuta", "Nororiente"],
], col_widths=[6, 3.5, 5])

doc.add_paragraph()

add_body("Impacto clave de la plataforma:", bold=True, color=DARK, size=11)
add_impact_card("⚡", "Estandarización Nacional", "Todos los centros utilizan la misma metodología de registro, conversión y reporte. Los datos son comparables entre regiones, sectores y periodos.", PRIMARY)
add_impact_card("🌍", "Huella de Carbono Medible", "Cálculo automático de emisiones TonCO2 para cada fuente energética, permitiendo cuantificar el impacto ambiental de cada empresa auditada.", SUCCESS)
add_impact_card("📊", "Inteligencia de Negocios", "Dashboards estratégico y nacional con visión consolidada de todos los proyectos, permitiendo decisiones basadas en datos a nivel de centro, región y país.", INFO)
add_impact_card("🔒", "Gobernanza y Seguridad", "Modelo de roles jerárquico con aislamiento de datos por centro. Cada universidad gestiona solo su información, con visión consolidada para la dirección nacional.", DANGER)
add_impact_card("📄", "Reportes Profesionales Automáticos", "Generación instantánea de informes PDF con gráficos, KPIs y análisis completo listos para entregar a la empresa auditada y a la UPME.", WARNING)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  ACCESO AL SISTEMA
# ══════════════════════════════════════════════════════════════════

add_section_header(3, "Acceso y Autenticación", "Diseño moderno con identidad visual del programa PEVI")

add_body("La pantalla de inicio de sesión presenta un diseño profesional que refleja la identidad del programa. Un panel lateral muestra los logos de las 5 universidades participantes sobre un fondo energético, mientras que el panel principal ofrece autenticación por usuario o correo electrónico.")

add_screenshot_placeholder("Pantalla de Login — Diseño a dos paneles con logos universitarios", height_cm=10)

add_body("Características de seguridad del acceso:", bold=True, color=DARK, size=10)
add_bullet("Autenticación dual: usuario tradicional o correo electrónico", "Doble vía de acceso:")
add_bullet("Restablecimiento por email con enlace seguro", "Recuperación de contraseña:")
add_bullet("4 niveles: Estudiante → Profesor → Director Centro → Director Nacional", "Roles jerárquicos:")
add_bullet("Cada centro solo ve y gestiona sus propios datos", "Aislamiento por centro:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  DASHBOARD
# ══════════════════════════════════════════════════════════════════

add_section_header(4, "Dashboard Inteligente", "Vista ejecutiva con indicadores clave en tiempo real")

add_body("Al ingresar, cada usuario recibe un dashboard personalizado según su rol. Los directores ven una vista ejecutiva con KPIs consolidados, mientras que profesores y estudiantes acceden a un panel operativo orientado a la acción.")

add_screenshot_placeholder("Dashboard Principal — Vista ejecutiva con 4 KPI cards y tabla de actividad reciente", height_cm=10)

add_body("Indicadores en tiempo real:", bold=True, color=DARK, size=10)

add_kpi_row([
    ("24", "PROYECTOS\nTOTALES", "e0f2fe"),
    ("8", "PROYECTOS\nACTIVOS", "fef3c7"),
    ("1.2M", "kWh\nAUDITADOS", "dcfce7"),
    ("→ BI", "MÉTRICAS\nESTRATÉGICAS", "f1f5f9"),
])

add_body("El dashboard incluye una tabla de actividad reciente con los últimos 10 proyectos actualizados, estado visual con badges de color, energía auditada y acceso directo a cada proyecto.")

add_body("Los profesores pueden iniciar nuevos proyectos directamente desde el dashboard con un solo clic, mientras que los estudiantes ven exclusivamente los proyectos donde están asignados como equipo de trabajo.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  GESTION DE PROYECTOS
# ══════════════════════════════════════════════════════════════════

add_section_header(5, "Gestión de Proyectos de Auditoría", "Motor central del sistema con flujo de estados, equipos y bitácora energética")

add_body("El módulo de proyectos es el corazón de la plataforma. Cada proyecto de auditoría atraviesa un ciclo de vida controlado (Borrador → Ejecución → Revisión → Finalizado) con asignación de equipos, registro de 6 fuentes energéticas, oportunidades de mejora y generación de informes.")

add_heading_h = doc.add_heading("Listado de Proyectos con Filtros Avanzados", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("El listado ofrece una vista completa con filtros por estado, fase, líder y centro. Los directores nacionales ven el portafolio de todo el país, mientras que cada centro tiene su propia vista filtrada.")

add_screenshot_placeholder("Lista de Proyectos — Filtros, badges de estado, equipo y acciones", height_cm=9)

add_heading_h = doc.add_heading("Creación de Proyectos con Selector de Equipo por Roles", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("El formulario de creación incluye un innovador selector de equipo que agrupa a los usuarios por rol (Estudiante, Profesor, Director). El usuario hace clic en el rol deseado, ve las personas disponibles como tarjetas interactivas, y los seleccionados aparecen como badges con opción de remover.")

add_screenshot_placeholder("Crear Proyecto — Formulario con selector de equipo por roles interactivo", height_cm=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  EXPEDIENTE DEL PROYECTO
# ══════════════════════════════════════════════════════════════════

add_section_header(6, "Expediente Integral del Proyecto", "Centro de comando con KPIs, gráficos, bitácora y análisis en una sola vista")

add_body("El expediente del proyecto es la vista más completa y poderosa del sistema. Concentra toda la información de una auditoría energética en una sola página interactiva: indicadores clave, 4 gráficos analíticos, bitácora de registros por fuente, oportunidades de mejora y documentación.")

add_screenshot_placeholder("Expediente del Proyecto — Barra de estado + 3 KPI cards superiores", height_cm=8)

add_body("3 KPIs principales en la cabecera:", bold=True, color=DARK, size=10)

add_kpi_row([
    ("245,800", "kWh EQUIVALENTES\nCONSUMO TOTAL", "e0f2fe"),
    ("85.3", "TonCO2/AÑO\nHUELLA DE CARBONO", "dcfce7"),
    ("$142.5M", "COP/AÑO\nCOSTO OPERATIVO", "fef3c7"),
])

add_heading_h = doc.add_heading("4 Gráficos Analíticos Interactivos", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("El expediente incluye 4 gráficos Chart.js interactivos que permiten visualizar la distribución energética, ambiental y económica de la empresa auditada:")

add_two_screenshots(
    "Gráfico 1: Matriz de Consumo\n(Dona — kWh por fuente)",
    "Gráfico 2: Huella de Carbono\n(Dona — TonCO2 por fuente)"
)
add_two_screenshots(
    "Gráfico 3: Balance Térmico\n(Barras — Eléctrico vs Térmico MBTU)",
    "Gráfico 4: Distribución de Costos\n(Barras — COP por fuente)"
)

doc.add_page_break()

add_heading_h = doc.add_heading("Bitácora de Registros Energéticos", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("7 tarjetas interactivas muestran el estado de cada fuente energética. Cada tarjeta resume los datos registrados (consumo, costo, emisiones) y ofrece acceso directo al formulario de registro.")

add_screenshot_placeholder("Bitácora — 7 tarjetas de fuentes energéticas (Producción + 6 combustibles)", height_cm=9)

add_styled_table([
    ["Fuente Energética", "Unidad", "Indicadores Mostrados"],
    ["⚡ Electricidad", "kWh", "Consumo kWh + Costo COP"],
    ["🔥 Gas Natural", "m³", "Volumen + Energía equivalente kWh"],
    ["⬛ Carbón Mineral", "Ton", "Masa + Emisiones TonCO2"],
    ["🔴 Fuel Oil", "Gal", "Volumen + Energía kWh"],
    ["🌿 Biomasa", "Ton", "Tipo de biomasa + Masa"],
    ["🔵 Gas Propano (GLP)", "kg", "Masa + Energía kWh"],
], col_widths=[4, 2, 9])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  REGISTRO ENERGETICO
# ══════════════════════════════════════════════════════════════════

add_section_header(7, "Registro Energético Inteligente", "Formularios estandarizados con conversión automática a kWh y cálculo de emisiones CO2")

add_body("El sistema de registro energético es uno de los componentes más potentes de la plataforma. Cada formulario está diseñado según los estándares de la UPME y realiza automáticamente la conversión a kWh equivalente y el cálculo de emisiones de CO2.")

add_screenshot_placeholder("Formulario de Registro Energético — Ejemplo con Gas Natural o Electricidad", height_cm=10)

add_body("Características del motor de cálculo:", bold=True, color=DARK, size=10)

add_impact_card("🔄", "Conversión Automática a kWh", "Cada combustible se convierte automáticamente a kWh equivalente usando la fórmula: Energía = (Consumo × Factor_Unidad × PC × Factor_Energía) / 3,600. El usuario solo ingresa el consumo original y el poder calorífico.", PRIMARY)
add_impact_card("🌱", "Cálculo de Emisiones CO2", "Las emisiones se calculan automáticamente: TonCO2 = (Consumo_Original × Factor_Emisión) / 1,000. El sistema usa factores de emisión oficiales referenciados por la UPME.", SUCCESS)
add_impact_card("🔢", "Manejo Inteligente de Números", "Los formularios aceptan números con comas (1,200.50), puntos decimales y diferentes formatos. El sistema limpia automáticamente los datos para evitar errores.", INFO)
add_impact_card("📐", "Conversor de Unidades Integrado", "Cada formulario incluye un conversor de unidades de poder calorífico (kJ, MJ, BTU, kcal) para facilitar el ingreso de datos desde diferentes fuentes técnicas.", WARNING)

add_screenshot_placeholder("Detalle del formulario — Secciones de consumo, costos, conversión y emisiones", height_cm=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  OPORTUNIDADES DE MEJORA
# ══════════════════════════════════════════════════════════════════

add_section_header(8, "Análisis de Oportunidades de Mejora", "Herramienta de simulación de escenarios de reducción energética")

add_body("La sección de Oportunidades de Mejora transforma la plataforma de un sistema de registro a una herramienta de análisis prospectivo. Permite establecer metas de reducción porcentual para cada fuente energética y calcula automáticamente los ahorros potenciales en energía, costos y emisiones.")

add_screenshot_placeholder("Panel de Oportunidades — Comparación Actual vs Proyectado + 3 KPIs de ahorro", height_cm=8)

add_body("Panel de impacto visual:", bold=True, color=DARK, size=10)

add_kpi_row([
    ("36,870", "kWh AHORRADOS\nENERGÍA POTENCIAL", "dcfce7"),
    ("$21.3M", "COP AHORRADOS\nIMPACTO ECONÓMICO", "fef3c7"),
    ("12.8", "TonCO2 EVITADAS\nIMPACTO AMBIENTAL", "e0f2fe"),
])

add_screenshot_placeholder("Gráfico de barras horizontales — Reducción por fuente energética + Tabla detallada", height_cm=8)

add_heading_h = doc.add_heading("OPMs: Proyectos de Mejora con Análisis Financiero", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Además de las metas porcentuales, el sistema permite registrar Oportunidades de Mejora concretas (OPMs) con análisis técnico-financiero completo: ahorro energético, costos evitados, inversión requerida, VPN, TIR y período de recuperación (payback).")

add_screenshot_placeholder("Tabla de OPMs — Proyectos identificados con métricas financieras y totales", height_cm=7)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  INFORME PDF
# ══════════════════════════════════════════════════════════════════

add_section_header(9, "Generación de Informes PDF", "Reportes profesionales generados con un solo clic, listos para entrega")

add_quote("Un informe profesional completo, generado en segundos, listo para entregar a la empresa auditada y a la UPME.")

add_body("El generador de informes PDF es una de las funcionalidades más valoradas de la plataforma. Con un solo clic desde el expediente del proyecto, el sistema genera un reporte profesional completo con toda la información de la auditoría energética.")

add_body("Contenido del informe generado:", bold=True, color=DARK, size=11)

table = doc.add_table(rows=8, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(table)
sections_pdf = [
    ("01  ENCABEZADO INSTITUCIONAL", "Nombre del proyecto, empresa auditada, centro PEVI, fecha de generación, estado y fase del proyecto."),
    ("02  RESUMEN EJECUTIVO (KPIs)", "Energía total (kWh), Emisiones (TonCO2), Costo operativo (COP), Indicador IDES de desempeño energético."),
    ("03  TABLA DE CONSUMOS", "Todas las fuentes con consumo original, kWh equivalente, emisiones y costos anuales. Fila de totales."),
    ("04  DISTRIBUCIÓN ENERGÉTICA", "Gráfico de barras horizontales por fuente con porcentajes y valores absolutos codificados por color."),
    ("05  BALANCE TÉRMICO", "Análisis eléctrico vs térmico en MBTU, costo por kWh, emisiones por unidad de producción."),
    ("06  OPORTUNIDADES DE MEJORA", "Ahorros proyectados en kWh, COP y TonCO2. Tabla de OPMs con métricas financieras."),
    ("07  EQUIPO DE TRABAJO", "Lista de miembros del proyecto con roles asignados."),
    ("08  METADATOS", "Información de generación, clasificación y referencias."),
]
for i, (titulo, desc) in enumerate(sections_pdf):
    c0 = table.rows[i].cells[0]
    c1 = table.rows[i].cells[1]
    c0.width = Cm(5)
    c1.width = Cm(10)

    if i % 2 == 0:
        set_cell_shading(c0, "f1f5f9")
        set_cell_shading(c1, "f1f5f9")

    p0 = c0.paragraphs[0]
    r0 = p0.add_run(titulo)
    r0.font.size = Pt(9)
    r0.font.bold = True
    r0.font.color.rgb = PRIMARY

    p1 = c1.paragraphs[0]
    r1 = p1.add_run(desc)
    r1.font.size = Pt(9)
    r1.font.color.rgb = GRAY_600

doc.add_paragraph()

add_screenshot_placeholder("Informe PDF generado — Primera página (Encabezado + Resumen Ejecutivo)", height_cm=10)

doc.add_page_break()

add_screenshot_placeholder("Informe PDF — Tabla de consumos energéticos detallada", height_cm=10)

add_screenshot_placeholder("Informe PDF — Gráfico de distribución energética por fuente", height_cm=10)

doc.add_page_break()

add_screenshot_placeholder("Informe PDF — Oportunidades de mejora y tabla de OPMs", height_cm=10)

add_body("El informe se descarga automáticamente como PDF con el nombre: Informe_PEVI_[Empresa]_[ID].pdf. Está diseñado para ser presentado directamente a la empresa auditada, a los directores del programa y a la UPME sin necesidad de edición adicional.", bold=False, size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  DASHBOARDS BI
# ══════════════════════════════════════════════════════════════════

add_section_header(10, "Inteligencia de Negocios", "Dashboards estratégico y nacional para la toma de decisiones basada en datos")

add_body("La plataforma incluye dos poderosos dashboards de Business Intelligence que consolidan la información de todos los proyectos y permiten análisis a nivel de centro, región y país.")

add_heading_h = doc.add_heading("Dashboard Estratégico", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Disponible para Directores de Centro y Director Nacional. Ofrece una visión consolidada de KPIs, gráficos de distribución energética y tabla detallada de todos los proyectos del alcance del usuario.")

add_screenshot_placeholder("Dashboard Estratégico — KPIs + Gráficos de dona y barras + Tabla de proyectos", height_cm=10)

add_heading_h = doc.add_heading("Dashboard Nacional", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Exclusivo para el Director Nacional. Presenta una vista comparativa de todos los centros PEVI del país, con ranking por energía auditada, emisiones y promedio por proyecto. Permite seleccionar un centro específico para ver su análisis detallado.")

add_screenshot_placeholder("Dashboard Nacional — Tabla comparativa de centros + KPIs nacionales", height_cm=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  GESTION DE EMPRESAS Y USUARIOS
# ══════════════════════════════════════════════════════════════════

add_section_header(11, "Gestión de Empresas y Usuarios", "Administración completa del ecosistema de actores del programa")

add_heading_h = doc.add_heading("Directorio de Empresas", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Registro centralizado de todas las empresas industriales auditadas, con datos de identificación, sector productivo, ubicación y contacto. Cada empresa está vinculada a su centro PEVI correspondiente.")

add_screenshot_placeholder("Lista de Empresas — Directorio con búsqueda y acciones", height_cm=7)

add_heading_h = doc.add_heading("Gestión de Equipo", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Panel de administración de usuarios con avatares, roles codificados por color, estado de actividad y registro de último acceso. Los directores pueden crear, editar y desactivar cuentas respetando la jerarquía de roles.")

add_screenshot_placeholder("Lista de Usuarios — Avatares, roles, estado y último acceso", height_cm=7)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  SITIO PUBLICO
# ══════════════════════════════════════════════════════════════════

add_section_header(12, "Portal Público de Transparencia", "Resultados agregados del programa accesibles sin autenticación")

add_body("La plataforma incluye un sitio público completo que muestra los resultados agregados y anonimizados del programa PEVI, promoviendo la transparencia y visibilidad del impacto nacional.")

add_heading_h = doc.add_heading("Páginas Públicas", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_styled_table([
    ["Página", "Contenido Destacado"],
    ["🏠 Inicio", "Landing page profesional con KPIs públicos y noticias recientes del programa"],
    ["👥 Nosotros", "Información institucional sobre el programa PEVI y sus objetivos"],
    ["🏛️ Centros", "Directorio interactivo de centros con modals que muestran gráficos por centro"],
    ["📚 Biblioteca", "Guías técnicas de UPME (bombeo, motores, vapor, etc.) con enlaces a PDFs"],
    ["📊 Resultados", "Visualizaciones interactivas con filtros, KPIs y exportación de datos"],
], col_widths=[3, 12])

doc.add_paragraph()

add_screenshot_placeholder("Sitio Público — Landing page / Home", height_cm=9)

add_screenshot_placeholder("Página de Resultados — Filtros + KPIs + Gráficos interactivos", height_cm=9)

doc.add_page_break()

add_heading_h = doc.add_heading("Página de Centros con Modals Interactivos", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Al seleccionar un centro, se abre un modal interactivo con gráficos Chart.js que muestran la distribución de proyectos por estado, los sectores industriales atendidos y la evolución de proyectos en el tiempo.")

add_two_screenshots(
    "Modal de Centro — Distribución\nde proyectos por estado",
    "Modal de Centro — Sectores\nindustriales atendidos"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  SEGURIDAD Y ARQUITECTURA
# ══════════════════════════════════════════════════════════════════

add_section_header(13, "Seguridad y Arquitectura Técnica", "Infraestructura robusta con múltiples capas de protección")

add_body("La Plataforma PEVI implementa un modelo de seguridad empresarial con múltiples capas de protección, garantizando la integridad, confidencialidad y disponibilidad de los datos energéticos nacionales.")

add_heading_h = doc.add_heading("Capas de Seguridad", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_styled_table([
    ["Capa", "Mecanismo", "Protección"],
    ["1. Autenticación", "Login por usuario o email + contraseña segura", "Solo usuarios autorizados acceden al sistema"],
    ["2. Autorización por Rol", "Decoradores @solo_directivos, @solo_lideres, @acceso_staff", "Cada rol accede solo a lo que le corresponde"],
    ["3. Aislamiento por Centro", "Filtrado automático en queries de BD", "Los centros no ven datos de otros centros"],
    ["4. Propiedad de Proyecto", "verificar_acceso_proyecto() en cada operación", "Solo el líder, equipo y director acceden al proyecto"],
    ["5. Anti-escalada", "Validación en formularios y servidor", "Directores no pueden crear usuarios de mayor rango"],
    ["6. CSP Headers", "Content-Security-Policy en middleware", "Previene ataques XSS e inyección de scripts"],
    ["7. Auditoría", "Logs de acceso, actividad y seguridad", "Trazabilidad completa de todas las acciones"],
], col_widths=[2.5, 5, 6.5])

doc.add_paragraph()

add_heading_h = doc.add_heading("Stack Tecnológico", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_styled_table([
    ["Componente", "Tecnología", "Propósito"],
    ["Backend", "Django 5.2 (Python)", "Framework web robusto con ORM y seguridad integrada"],
    ["Base de Datos", "PostgreSQL", "BD relacional de alto rendimiento"],
    ["Frontend", "Bootstrap 5 + Chart.js", "Interfaz responsiva con gráficos interactivos"],
    ["Generación PDF", "WeasyPrint", "Renderizado HTML→PDF de alta calidad"],
    ["Infraestructura", "AWS EC2 + S3", "Servidor en la nube con almacenamiento escalable"],
    ["Dominio", "pevicolombia.com", "URL profesional con certificado SSL"],
], col_widths=[3, 4, 8])

doc.add_paragraph()

add_heading_h = doc.add_heading("Sistema de Logs y Auditoría", level=2)
for r in add_heading_h.runs: r.font.color.rgb = DARK

add_body("Toda acción en la plataforma queda registrada en un sistema de logs multinivel:")
add_bullet("Registra inicios de sesión, cierres y intentos fallidos", "Log de Acceso:")
add_bullet("Registra creación, edición y eliminación de datos", "Log de Actividad:")
add_bullet("Registra accesos denegados y bloqueos de escalada de privilegios", "Log de Seguridad:")
add_bullet("Registra excepciones y errores del sistema", "Log de Errores:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  CONCLUSIONES
# ══════════════════════════════════════════════════════════════════

add_section_header(14, "Conclusiones", "Una plataforma a la altura del reto energético nacional")

add_quote("La Plataforma PEVI demuestra que es posible construir un sistema nacional de gestión energética que sea potente, usable y seguro al mismo tiempo.")

add_body("La Plataforma PEVI representa un avance significativo en la digitalización de la gestión energética industrial en Colombia. Sus capacidades clave incluyen:", size=10)

doc.add_paragraph()

add_impact_card("✅", "Cobertura Nacional Integrada", "5 centros universitarios conectados en una sola plataforma, con datos estandarizados y comparables a nivel nacional.", SUCCESS)
add_impact_card("✅", "Motor de Cálculo Automático", "Conversión transparente de 6 fuentes energéticas a kWh equivalente y cálculo automático de emisiones CO2, eliminando errores manuales.", SUCCESS)
add_impact_card("✅", "Análisis Prospectivo", "No solo registra el presente: permite simular escenarios de reducción y evaluar el impacto financiero y ambiental de las mejoras propuestas.", SUCCESS)
add_impact_card("✅", "Informes Profesionales Instantáneos", "Generación automática de reportes PDF completos con gráficos, KPIs y análisis, listos para entrega sin edición adicional.", SUCCESS)
add_impact_card("✅", "Inteligencia para Decisiones", "Dashboards de BI a nivel estratégico y nacional que transforman datos dispersos en información accionable para tomadores de decisión.", SUCCESS)
add_impact_card("✅", "Transparencia Pública", "Portal abierto con resultados agregados y anonimizados que promueve la rendición de cuentas del programa.", SUCCESS)
add_impact_card("✅", "Seguridad Empresarial", "7 capas de protección con aislamiento por centro, roles jerárquicos, auditoría completa y protección contra ataques web.", SUCCESS)

doc.add_paragraph()

add_body("La plataforma no solo cumple con los requisitos operativos del programa PEVI, sino que establece un estándar de referencia para la gestión digital de auditorías energéticas en América Latina.", bold=True, color=DARK, size=10)

# ══════════════════════════════════════════════════════════════════
#  CONTRAPORTADA
# ══════════════════════════════════════════════════════════════════

doc.add_page_break()

for _ in range(7):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PLATAFORMA PEVI")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Programa de Evaluación de Eficiencia Energética Industrial")
run.font.size = Pt(11)
run.font.color.rgb = GRAY_500

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("https://pevicolombia.com")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = PRIMARY

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UPME  •  Ministerio de Energía  •  ONUDI")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_600

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Colombia, 2026")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_400

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("© Todos los derechos reservados")
run.font.size = Pt(9)
run.font.color.rgb = GRAY_400

# ══════════════════════════════════════════════════════════════════
#  GUARDAR
# ══════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Informe_Impacto_PEVI.docx")
doc.save(output_path)
print(f"✓ Informe generado: {output_path}")
print(f"  Tamaño: {os.path.getsize(output_path) / 1024:.0f} KB")

# Contar placeholders
count = sum(1 for s in sections_pdf) # rough count
print(f"  Espacios para capturas: ~25 recuadros marcados con 📷")
